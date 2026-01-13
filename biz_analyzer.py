"""
BizBuySell Business Opportunity Analyzer
A Streamlit app for analyzing business-for-sale listings.
"""

import os
import re
import io
import streamlit as st
import pandas as pd
from dotenv import load_dotenv
from scraper import scrape_listings, is_direct_listing_url, scrape_single_listing, SkippedListing
from analyzer import analyze_businesses, AnalysisResult

# Try to import python-docx for Word export
try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Load environment variables from .env file
load_dotenv()


def clean_ai_text(text: str) -> str:
    """Remove formatting characters and escape special characters in AI-generated text."""
    if not text:
        return text
    # Remove triple backticks but KEEP content (with optional language identifier on first line)
    text = re.sub(r'```[a-zA-Z]*\n?', '', text)  # Opening ```lang or ```
    text = re.sub(r'```', '', text)  # Closing ```
    # Remove double backticks but keep content
    text = re.sub(r'``', '', text)
    # Remove single backticks
    text = text.replace('`', '')
    # Also handle unicode variants of backticks
    text = text.replace('\u0060', '')  # grave accent
    text = text.replace('\u2018', '')  # left single quote
    text = text.replace('\u2019', '')  # right single quote
    # Escape dollar signs to prevent Streamlit LaTeX rendering
    text = text.replace('$', '\\$')
    return text


def clean_text_for_doc(text: str) -> str:
    """Clean text for Word document (no escaping needed)."""
    if not text:
        return text
    # Remove backticks
    text = re.sub(r'```[a-zA-Z]*\n?', '', text)
    text = re.sub(r'```', '', text)
    text = re.sub(r'``', '', text)
    text = text.replace('`', '')
    text = text.replace('\u0060', '')
    text = text.replace('\u2018', "'")
    text = text.replace('\u2019', "'")
    return text


def sanitize_for_csv(text: str) -> str:
    """Sanitize text for CSV export - replace special characters with ASCII equivalents."""
    if not text:
        return text
    # Replace emojis with text equivalents
    replacements = {
        '🌟': '[HIGHLY RECOMMENDED]',
        '✅': '[RECOMMENDED]',
        '⚠️': '[CAUTION]',
        '⚡': '[HIGH RISK]',
        '❌': '[NOT RECOMMENDED]',
        '💪': '',
        '🚀': '',
        '🔴': '',
        '📊': '',
        '💰': '',
        '📈': '',
        '🎯': '',
        '🏆': '',
        '📍': '',
        '📝': '',
        '🔍': '',
        '🚩': '',
        '💡': '',
        '📋': '',
        '🤖': '',
        '—': '-',  # em dash
        '–': '-',  # en dash
        '"': '"',  # smart quotes
        '"': '"',
        ''': "'",
        ''': "'",
        '…': '...',
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    # Remove any remaining non-ASCII characters
    text = text.encode('ascii', 'ignore').decode('ascii')
    return text.strip()


def generate_word_report(results: list) -> bytes:
    """Generate a Word document with the full analysis report."""
    if not DOCX_AVAILABLE:
        return None

    doc = Document()

    # Title
    title = doc.add_heading('Business Opportunity Analysis Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Summary section
    doc.add_heading('Executive Summary', level=1)
    doc.add_paragraph(f"Total businesses analyzed: {len(results)}")
    if results:
        avg_score = sum(r.overall_score for r in results) / len(results)
        doc.add_paragraph(f"Average score: {avg_score:.1f}/100")

        # Top 3 Recommendations
        doc.add_heading('Top 3 Recommendations', level=2)
        top_3 = results[:3]
        medals = ["1st", "2nd", "3rd"]

        for i, r in enumerate(top_3):
            doc.add_heading(f"{medals[i]} Place: {r.business.name}", level=3)
            doc.add_paragraph(f"Score: {r.overall_score:.0f}/100")
            doc.add_paragraph(f"Asking Price: {format_currency(r.business.asking_price)}")
            doc.add_paragraph(f"Cash Flow: {format_currency(r.business.cash_flow)}")
            if r.metrics.roi_percent:
                doc.add_paragraph(f"ROI: {r.metrics.roi_percent:.1f}%")

            # Top strength
            if r.swot.strengths:
                doc.add_paragraph(f"Top Strength: {clean_text_for_doc(r.swot.strengths[0])}")

            # Main risk
            if r.swot.weaknesses:
                doc.add_paragraph(f"Main Risk: {clean_text_for_doc(r.swot.weaknesses[0])}")
            elif r.viability and r.viability.key_risks:
                doc.add_paragraph(f"Main Risk: {clean_text_for_doc(r.viability.key_risks[0])}")

            # Recommendation summary
            doc.add_paragraph(f"Verdict: {clean_text_for_doc(r.recommendation)}")
            doc.add_paragraph(f"Listing URL: {r.business.url}")
            doc.add_paragraph("")  # Spacing

    doc.add_page_break()

    # Individual business analyses
    for r in results:
        # Business header
        doc.add_heading(f"#{r.rank} - {r.business.name}", level=1)
        doc.add_paragraph(f"Overall Score: {r.overall_score:.0f}/100")
        doc.add_paragraph(f"URL: {r.business.url}")

        # Details section
        doc.add_heading('Details', level=2)
        doc.add_paragraph(f"Location: {r.business.location or 'Not specified'}")
        doc.add_paragraph(f"Category: {r.business.category or 'Not specified'}")
        if r.business.employees:
            doc.add_paragraph(f"Employees: {r.business.employees}")
        if r.business.year_established:
            doc.add_paragraph(f"Established: {r.business.year_established}")

        if r.business.description:
            doc.add_heading('Description', level=3)
            doc.add_paragraph(clean_text_for_doc(r.business.description))

        # Business Metrics
        doc.add_heading('Business Metrics', level=2)
        doc.add_paragraph(f"Asking Price: {format_currency(r.business.asking_price)}")
        doc.add_paragraph(f"Cash Flow: {format_currency(r.business.cash_flow)}")
        if r.business.gross_revenue:
            doc.add_paragraph(f"Gross Revenue: {format_currency(r.business.gross_revenue)}")

        metrics_dict = r.metrics.to_dict()
        for key, value in metrics_dict.items():
            doc.add_paragraph(f"{key}: {value}")

        # SWOT Analysis
        doc.add_heading('SWOT Analysis', level=2)

        if r.swot.strengths:
            doc.add_heading('Strengths', level=3)
            for s in r.swot.strengths:
                doc.add_paragraph(f"• {clean_text_for_doc(s)}", style='List Bullet')

        if r.swot.weaknesses:
            doc.add_heading('Weaknesses', level=3)
            for w in r.swot.weaknesses:
                doc.add_paragraph(f"• {clean_text_for_doc(w)}", style='List Bullet')

        if r.swot.opportunities:
            doc.add_heading('Opportunities', level=3)
            for o in r.swot.opportunities:
                doc.add_paragraph(f"• {clean_text_for_doc(o)}", style='List Bullet')

        if r.swot.threats:
            doc.add_heading('Threats', level=3)
            for t in r.swot.threats:
                doc.add_paragraph(f"• {clean_text_for_doc(t)}", style='List Bullet')

        # Competitive Analysis
        if r.competitive:
            doc.add_heading('Competitive Analysis', level=2)
            if r.competitive.market_position:
                doc.add_paragraph(f"Market Position: {clean_text_for_doc(r.competitive.market_position)}")
            if r.competitive.valuation_assessment:
                doc.add_paragraph(f"Valuation Assessment: {clean_text_for_doc(r.competitive.valuation_assessment)}")

            if r.competitive.competitive_advantages:
                doc.add_heading('Competitive Advantages', level=3)
                for adv in r.competitive.competitive_advantages:
                    doc.add_paragraph(f"• {clean_text_for_doc(adv)}", style='List Bullet')

            if r.competitive.competitive_threats:
                doc.add_heading('Competitive Threats', level=3)
                for threat in r.competitive.competitive_threats:
                    doc.add_paragraph(f"• {clean_text_for_doc(threat)}", style='List Bullet')

        # Market Analysis
        if r.market_analysis:
            doc.add_heading('Market Analysis', level=2)
            doc.add_paragraph(clean_text_for_doc(r.market_analysis))

        # Viability Assessment
        if r.viability:
            doc.add_heading('Viability Assessment', level=2)
            viable_status = "Viable" if r.viability.is_viable else "Not Viable"
            doc.add_paragraph(f"Status: {viable_status} (Confidence: {r.viability.confidence})")
            doc.add_paragraph(f"Viability Score: {r.viability.viability_score}/100")

            if r.viability.verdict:
                doc.add_paragraph(f"Verdict: {clean_text_for_doc(r.viability.verdict)}")

            if r.viability.red_flags:
                doc.add_heading('Red Flags', level=3)
                for flag in r.viability.red_flags:
                    doc.add_paragraph(f"• {clean_text_for_doc(flag)}", style='List Bullet')

            if r.viability.key_risks:
                doc.add_heading('Key Risks', level=3)
                for risk in r.viability.key_risks:
                    doc.add_paragraph(f"• {clean_text_for_doc(risk)}", style='List Bullet')

            if r.viability.key_opportunities:
                doc.add_heading('Key Opportunities', level=3)
                for opp in r.viability.key_opportunities:
                    doc.add_paragraph(f"• {clean_text_for_doc(opp)}", style='List Bullet')

            if r.viability.due_diligence_items:
                doc.add_heading('Due Diligence Items', level=3)
                for item in r.viability.due_diligence_items:
                    doc.add_paragraph(f"• {clean_text_for_doc(item)}", style='List Bullet')

        # Recommendation
        doc.add_heading('Recommendation', level=2)
        doc.add_paragraph(clean_text_for_doc(r.recommendation))

        # Page break between businesses (except for last one)
        if r.rank < len(results):
            doc.add_page_break()

    # Save to bytes
    doc_bytes = io.BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)
    return doc_bytes.getvalue()


# Page configuration
st.set_page_config(
    page_title="Business Opportunity Analyzer",
    page_icon="🏢",
    layout="wide",
    initial_sidebar_state="collapsed"
)

# Custom CSS for a distinctive, modern look
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Space+Grotesk:wght@300;400;500;600;700&family=JetBrains+Mono:wght@400;500&display=swap');

    :root {
        --primary: #0f172a;
        --secondary: #1e293b;
        --accent: #f59e0b;
        --accent-light: #fbbf24;
        --success: #10b981;
        --warning: #f59e0b;
        --danger: #ef4444;
        --text: #f8fafc;
        --text-muted: #94a3b8;
        --card-bg: #1e293b;
        --gradient-1: linear-gradient(135deg, #0f172a 0%, #1e293b 50%, #334155 100%);
    }

    .stApp {
        background: var(--gradient-1);
    }

    .main .block-container {
        padding-top: 2rem;
        padding-bottom: 2rem;
        max-width: 1400px;
    }

    h1, h2, h3, h4, h5, h6 {
        font-family: 'Space Grotesk', sans-serif !important;
        color: #f8fafc !important;
    }

    /* Force readable text colors on all elements */
    p, li, span, div, label {
        color: #e2e8f0 !important;
    }

    /* Markdown text */
    .stMarkdown, .stMarkdown p, .stMarkdown li, .stMarkdown span {
        color: #e2e8f0 !important;
    }

    /* Bold text should be brighter */
    strong, b {
        color: #ffffff !important;
        font-weight: 600 !important;
    }

    /* Links */
    a {
        color: #fbbf24 !important;
    }

    /* Expander content */
    .streamlit-expanderContent {
        color: #e2e8f0 !important;
    }

    .streamlit-expanderContent p,
    .streamlit-expanderContent li,
    .streamlit-expanderContent span {
        color: #e2e8f0 !important;
    }

    .hero-title {
        font-family: 'Space Grotesk', sans-serif;
        font-size: 3rem;
        font-weight: 700;
        background: linear-gradient(135deg, #f59e0b 0%, #fbbf24 50%, #fcd34d 100%);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        background-clip: text;
        margin-bottom: 0.5rem;
        text-align: center;
    }

    .hero-subtitle {
        font-family: 'Space Grotesk', sans-serif;
        font-size: 1.2rem;
        color: var(--text-muted);
        text-align: center;
        margin-bottom: 2rem;
    }

    .metric-card {
        background: linear-gradient(145deg, #1e293b 0%, #334155 100%);
        border: 1px solid #475569;
        border-radius: 16px;
        padding: 1.5rem;
        margin: 0.5rem 0;
        box-shadow: 0 4px 20px rgba(0, 0, 0, 0.3);
    }

    .metric-value {
        font-family: 'JetBrains Mono', monospace;
        font-size: 2rem;
        font-weight: 600;
        color: var(--accent);
    }

    .metric-label {
        font-family: 'Space Grotesk', sans-serif;
        font-size: 0.9rem;
        color: var(--text-muted);
        text-transform: uppercase;
        letter-spacing: 0.05em;
    }

    .rank-badge {
        display: inline-flex;
        align-items: center;
        justify-content: center;
        width: 40px;
        height: 40px;
        border-radius: 50%;
        font-family: 'JetBrains Mono', monospace;
        font-weight: 700;
        font-size: 1.2rem;
    }

    .rank-1 { background: linear-gradient(135deg, #f59e0b, #fcd34d); color: #0f172a; }
    .rank-2 { background: linear-gradient(135deg, #94a3b8, #cbd5e1); color: #0f172a; }
    .rank-3 { background: linear-gradient(135deg, #b45309, #d97706); color: #fff; }

    .swot-section {
        background: var(--card-bg);
        border-radius: 12px;
        padding: 1rem;
        margin: 0.5rem 0;
    }

    .swot-strength { border-left: 4px solid #10b981; }
    .swot-weakness { border-left: 4px solid #ef4444; }
    .swot-opportunity { border-left: 4px solid #3b82f6; }
    .swot-threat { border-left: 4px solid #f59e0b; }

    .score-excellent { color: #10b981; }
    .score-good { color: #3b82f6; }
    .score-moderate { color: #f59e0b; }
    .score-poor { color: #ef4444; }

    .top-pick-card {
        background: linear-gradient(145deg, #1e293b 0%, #0f172a 100%);
        border: 2px solid #f59e0b;
        border-radius: 20px;
        padding: 2rem;
        margin: 1rem 0;
        box-shadow: 0 8px 32px rgba(245, 158, 11, 0.2);
    }

    .stButton > button, .stLinkButton > a, .stDownloadButton > button {
        font-family: 'Space Grotesk', sans-serif;
        font-weight: 600;
        background: linear-gradient(135deg, #f59e0b 0%, #d97706 100%) !important;
        color: #0f172a !important;
        border: none;
        border-radius: 12px;
        padding: 0.75rem 2rem;
        font-size: 1.1rem;
        transition: all 0.3s ease;
        width: 100%;
    }

    .stButton > button:hover, .stLinkButton > a:hover, .stDownloadButton > button:hover {
        background: linear-gradient(135deg, #fbbf24 0%, #f59e0b 100%) !important;
        box-shadow: 0 4px 20px rgba(245, 158, 11, 0.4);
        transform: translateY(-2px);
        color: #0f172a !important;
    }

    /* Link button specific */
    .stLinkButton a, [data-testid="stLinkButton"] a {
        color: #0f172a !important;
        text-decoration: none !important;
    }

    /* Download button specific */
    .stDownloadButton button, [data-testid="stDownloadButton"] button {
        color: #0f172a !important;
    }

    .stTextInput > div > div > input {
        font-family: 'JetBrains Mono', monospace;
        background: var(--secondary);
        border: 2px solid #475569;
        border-radius: 12px;
        color: var(--text);
        padding: 0.75rem 1rem;
    }

    .stTextInput > div > div > input:focus {
        border-color: var(--accent);
        box-shadow: 0 0 0 2px rgba(245, 158, 11, 0.2);
    }

    .stNumberInput > div > div > input {
        font-family: 'JetBrains Mono', monospace;
        background: var(--secondary);
        border: 2px solid #475569;
        border-radius: 12px;
        color: var(--text);
    }

    .stDataFrame {
        border-radius: 12px;
        overflow: hidden;
    }

    div[data-testid="stDataFrame"] > div {
        background: #1e293b;
        border-radius: 12px;
    }

    /* Table text colors */
    .stDataFrame td, .stDataFrame th {
        color: #e2e8f0 !important;
    }

    /* DataFrame cell text */
    [data-testid="stDataFrame"] * {
        color: #e2e8f0 !important;
    }

    .stProgress > div > div > div {
        background: linear-gradient(90deg, #f59e0b, #fbbf24);
    }

    /* Expander styling */
    [data-testid="stExpander"] {
        background: #1e293b;
        border: 1px solid #475569;
        border-radius: 12px;
    }

    /* Expander header text color */
    [data-testid="stExpander"] [data-testid="stMarkdownContainer"] p {
        color: #ffffff !important;
    }

    /* Info/Success/Warning boxes */
    .stAlert, [data-testid="stAlert"] {
        color: #1e293b !important;
    }

    /* Captions and small text */
    .stCaption, small, .caption {
        color: #94a3b8 !important;
    }

    /* Metric labels */
    [data-testid="stMetricLabel"] {
        color: #cbd5e1 !important;
    }

    [data-testid="stMetricValue"] {
        color: #ffffff !important;
    }

    .info-badge {
        display: inline-block;
        padding: 0.25rem 0.75rem;
        border-radius: 20px;
        font-size: 0.85rem;
        font-weight: 500;
        margin-right: 0.5rem;
    }

    .badge-price { background: #312e81; color: #a5b4fc; }
    .badge-roi { background: #064e3b; color: #6ee7b7; }
    .badge-score { background: #78350f; color: #fcd34d; }
</style>
""", unsafe_allow_html=True)


def format_currency(value: float) -> str:
    """Format a number as currency."""
    if value is None:
        return "N/A"
    if value >= 1_000_000:
        return f"${value/1_000_000:.1f}M"
    elif value >= 1_000:
        return f"${value/1_000:.0f}K"
    else:
        return f"${value:.0f}"


def render_header():
    """Render the app header."""
    st.markdown('<h1 class="hero-title">🏢 Business Opportunity Analyzer</h1>', unsafe_allow_html=True)
    st.markdown('<p class="hero-subtitle">AI-powered analysis of BizBuySell listings • SWOT • Metrics • Rankings</p>', unsafe_allow_html=True)



def render_input_form():
    """Render the input form and return values."""
    col1, col2 = st.columns([3, 1])

    with col1:
        url = st.text_input(
            "🔗 BizBuySell URL",
            placeholder="https://www.bizbuysell.com/california/san-francisco-bay-area-businesses-for-sale/... or .../business-opportunity/...",
            help="Paste a BizBuySell search page URL or a direct business listing URL"
        )

    with col2:
        limit = st.number_input(
            "📊 Listings to Analyze",
            min_value=1,
            max_value=100,
            value=5,
            help="Number of business listings to analyze (1-100)"
        )

    # Load ScraperAPI key from environment
    scraper_api_key = os.getenv("SCRAPER_API_KEY", "")

    analyze_clicked = st.button("🚀 Analyze Opportunities", use_container_width=True)

    # Show timing estimate and filtering info
    openai_key = os.getenv("OPENAI_API_KEY", "")
    if openai_key and not openai_key.startswith("your_"):
        st.caption(f"⏱️ Estimated time: ~{limit * 10 + 10}s (includes AI filtering & analysis)")
        st.caption("🔍 Auto-filtering: dental, legal, CPA, medical, and restaurants are excluded")

    return url, limit, analyze_clicked, scraper_api_key if scraper_api_key else None


def render_comparison_table(results: list[AnalysisResult]):
    """Render the comparison table."""
    st.markdown("---")
    st.markdown("## 📊 Comparison Table")

    # Build DataFrame
    table_data = []
    for r in results:
        # Get AI viability score if available
        ai_score = f"{r.viability.viability_score}" if r.viability and r.viability.viability_score else "N/A"

        row = {
            "Rank": f"#{r.rank}",
            "Business": r.business.name,
            "Price": format_currency(r.business.asking_price),
            "Cash Flow": format_currency(r.business.cash_flow),
            "ROI": f"{r.metrics.roi_percent:.1f}%" if r.metrics.roi_percent else "N/A",
            "Rule Score": f"{r.rule_based_score:.0f}",
            "AI Score": ai_score,
            "Total Score": f"{r.overall_score:.0f}",
        }
        table_data.append(row)

    df = pd.DataFrame(table_data)

    # Style the dataframe
    st.dataframe(
        df,
        use_container_width=True,
        hide_index=True,
        column_config={
            "Rank": st.column_config.TextColumn("Rank", width="small"),
            "Business": st.column_config.TextColumn("Business", width="large"),
            "Price": st.column_config.TextColumn("Price", width="small"),
            "Cash Flow": st.column_config.TextColumn("Cash Flow", width="small"),
            "ROI": st.column_config.TextColumn("ROI", width="small"),
            "Rule Score": st.column_config.TextColumn("Rule Score", width="small"),
            "AI Score": st.column_config.TextColumn("AI Score", width="small"),
            "Total Score": st.column_config.TextColumn("Total Score", width="small"),
        }
    )

    # Scoring methodology note
    st.caption("**Scoring:** Rule Score = ROI + Payback + Age + SWOT + Data completeness (0-100). "
               "AI Score = OpenAI viability assessment (0-100). "
               "**Total Score = Rule Score × 0.6 + AI Score × 0.4**")

    # Expandable details for each business
    st.markdown("### 📋 Detailed Analysis")

    for r in results:
        ai_badge = " 🤖" if r.ai_enhanced else ""
        with st.expander(f"#{r.rank} - {r.business.name} | Score: {r.overall_score:.0f}/100{ai_badge}"):
            # Details (collapsible) - includes description
            with st.expander("📍 Details", expanded=False):
                st.markdown(f"**Location:** {r.business.location or 'Not specified'}")
                st.markdown(f"**Category:** {r.business.category or 'Not specified'}")
                if r.business.employees:
                    st.markdown(f"**Employees:** {r.business.employees}")
                if r.business.year_established:
                    st.markdown(f"**Established:** {r.business.year_established}")
                if r.business.description:
                    st.markdown("**📝 Description:**")
                    st.markdown(f"{clean_ai_text(r.business.description)}")

            # Business Metrics (collapsible)
            with st.expander("📈 Business Metrics", expanded=False):
                st.markdown(f"**Asking Price:** {format_currency(r.business.asking_price)}")
                st.markdown(f"**Cash Flow:** {format_currency(r.business.cash_flow)}")
                if r.business.gross_revenue:
                    st.markdown(f"**Gross Revenue:** {format_currency(r.business.gross_revenue)}")
                metrics_dict = r.metrics.to_dict()
                for key, value in metrics_dict.items():
                    st.markdown(f"**{key}:** {value}")

            # SWOT Analysis (collapsible)
            with st.expander("🎯 SWOT Analysis", expanded=False):
                if r.swot.strengths:
                    st.markdown("**💪 Strengths:**")
                    for s in r.swot.strengths[:3]:
                        st.markdown(f"- {clean_ai_text(s)}")

                if r.swot.weaknesses:
                    st.markdown("**⚠️ Weaknesses:**")
                    for w in r.swot.weaknesses[:3]:
                        st.markdown(f"- {clean_ai_text(w)}")

                if r.swot.opportunities:
                    st.markdown("**🚀 Opportunities:**")
                    for o in r.swot.opportunities[:3]:
                        st.markdown(f"- {clean_ai_text(o)}")

                if r.swot.threats:
                    st.markdown("**🔴 Threats:**")
                    for t in r.swot.threats[:3]:
                        st.markdown(f"- {clean_ai_text(t)}")

            # Competitive Analysis (AI-powered with Tavily market data)
            if r.competitive or r.market_data:
                # Competitive Analysis (collapsible)
                if r.competitive:
                    with st.expander("🏆 Competitive Analysis", expanded=False):
                        if r.competitive.market_position:
                            st.markdown(f"**📊 Market Position:** {clean_ai_text(r.competitive.market_position)}")

                        if r.competitive.valuation_assessment:
                            st.markdown(f"**💰 Valuation Assessment:** {clean_ai_text(r.competitive.valuation_assessment)}")

                        ccol1, ccol2 = st.columns(2)

                        with ccol1:
                            if r.competitive.competitive_advantages:
                                st.markdown("**✅ Competitive Advantages:**")
                                for adv in r.competitive.competitive_advantages[:4]:
                                    st.markdown(f"- {clean_ai_text(adv)}")

                        with ccol2:
                            if r.competitive.competitive_threats:
                                st.markdown("**⚠️ Competitive Threats:**")
                                for threat in r.competitive.competitive_threats[:4]:
                                    st.markdown(f"- {clean_ai_text(threat)}")

                # Market Analysis (collapsible)
                if r.market_analysis:
                    with st.expander("📈 Market Analysis", expanded=False):
                        st.markdown(clean_ai_text(r.market_analysis))
                elif r.market_data:
                    # Fallback to raw data if AI analysis not available
                    with st.expander("📈 Raw Market Research Data", expanded=False):
                        st.markdown(clean_ai_text(r.market_data))

            # Viability Assessment (AI-powered, collapsible)
            if r.viability:
                with st.expander("🔍 AI Viability Assessment", expanded=False):
                    # Viability status
                    viable_icon = "✅" if r.viability.is_viable else "❌"
                    st.markdown(f"**Status:** {viable_icon} {'Viable' if r.viability.is_viable else 'Not Viable'} (Confidence: {r.viability.confidence})")
                    st.markdown(f"**Viability Score:** {r.viability.viability_score}/100")

                    if r.viability.verdict:
                        st.info(f"**Verdict:** {clean_ai_text(r.viability.verdict)}")

                    vcol1, vcol2 = st.columns(2)

                    with vcol1:
                        if r.viability.red_flags:
                            st.markdown("**🚩 Red Flags:**")
                            for flag in r.viability.red_flags[:3]:
                                st.markdown(f"- ⚠️ {clean_ai_text(flag)}")

                        if r.viability.key_risks:
                            st.markdown("**⚡ Key Risks:**")
                            for risk in r.viability.key_risks[:3]:
                                st.markdown(f"- {clean_ai_text(risk)}")

                    with vcol2:
                        if r.viability.key_opportunities:
                            st.markdown("**💡 Key Opportunities:**")
                            for opp in r.viability.key_opportunities[:3]:
                                st.markdown(f"- {clean_ai_text(opp)}")

                        if r.viability.due_diligence_items:
                            st.markdown("**📋 Due Diligence:**")
                            for item in r.viability.due_diligence_items[:3]:
                                st.markdown(f"- {clean_ai_text(item)}")

            st.markdown("")
            st.link_button("🔗 View Full Listing on BizBuySell", r.business.url, use_container_width=True)


def render_top_picks(results: list[AnalysisResult]):
    """Render the top 3 recommendations."""
    st.markdown("---")
    st.markdown("## 🏆 Top 3 Recommendations")

    top_3 = results[:3]

    cols = st.columns(3)

    medals = ["🥇", "🥈", "🥉"]
    colors = ["#f59e0b", "#94a3b8", "#b45309"]

    for i, (col, result) in enumerate(zip(cols, top_3)):
        with col:
            st.markdown(f"""
            <div class="top-pick-card" style="border-color: {colors[i]};">
                <div style="text-align: center; margin-bottom: 1rem;">
                    <span style="font-size: 3rem;">{medals[i]}</span>
                    <h3 style="margin: 0.5rem 0; font-size: 1.1rem; color: #f8fafc;">
                        {result.business.name}
                    </h3>
                </div>
                <div style="text-align: center; margin-bottom: 1rem;">
                    <span class="metric-value" style="font-size: 2.5rem;">{result.overall_score:.0f}</span>
                    <span style="color: #94a3b8; font-size: 1rem;">/100</span>
                </div>
            </div>
            """, unsafe_allow_html=True)

            st.markdown(f"**💰 Price:** {format_currency(result.business.asking_price)}")
            st.markdown(f"**📈 ROI:** {result.metrics.roi_percent:.1f}%" if result.metrics.roi_percent else "**📈 ROI:** N/A")

            # Viability badge
            if result.viability:
                viable_icon = "✅" if result.viability.is_viable else "❌"
                st.markdown(f"**🔍 Viable:** {viable_icon} {result.viability.confidence}")

            st.markdown("---")
            st.markdown(f"**💪 Top Strength:**")
            if result.swot.strengths:
                st.markdown(f"_{result.swot.strengths[0][:80]}{'...' if len(result.swot.strengths[0]) > 80 else ''}_")
            else:
                st.markdown("_Not identified_")

            st.markdown(f"**⚠️ Main Risk:**")
            if result.swot.weaknesses:
                st.markdown(f"_{result.swot.weaknesses[0][:80]}{'...' if len(result.swot.weaknesses[0]) > 80 else ''}_")
            elif result.viability and result.viability.key_risks:
                st.markdown(f"_{result.viability.key_risks[0][:80]}_")
            else:
                st.markdown("_Not identified_")

            st.markdown("---")
            st.markdown(result.recommendation)

            st.link_button("🔗 View Listing", result.business.url, use_container_width=True)


def is_streamlit_cloud() -> bool:
    """Check if running on Streamlit Cloud platform."""
    # Streamlit Cloud runs from /mount/src/ directory
    return os.path.exists("/mount/src") or os.getenv("STREAMLIT_SHARING_MODE") == "true"


def check_password():
    """Returns True if the user has entered the correct password."""
    # Skip password check when running locally (not on Streamlit Cloud)
    if not is_streamlit_cloud():
        return True

    # Get password from environment variable
    correct_password = os.getenv("APP_PASSWORD", "")

    # If no password is set, allow access
    if not correct_password:
        return True

    # Initialize session state for authentication
    if "authenticated" not in st.session_state:
        st.session_state.authenticated = False

    # If already authenticated, return True
    if st.session_state.authenticated:
        return True

    # Show password input
    st.markdown('<h1 class="hero-title">🔐 Business Opportunity Analyzer</h1>', unsafe_allow_html=True)
    st.markdown('<p class="hero-subtitle">Please enter the password to access the app</p>', unsafe_allow_html=True)

    password = st.text_input("Password", type="password", key="password_input")

    if st.button("🔓 Login", use_container_width=True):
        if password == correct_password:
            st.session_state.authenticated = True
            st.rerun()
        else:
            st.error("❌ Incorrect password. Please try again.")

    return False


def main():
    """Main application entry point."""
    # Password protection
    if not check_password():
        return

    render_header()

    # Input form
    url, limit, analyze_clicked, scraper_api_key = render_input_form()

    # Session state for results
    if 'results' not in st.session_state:
        st.session_state.results = None
    if 'skipped_listings' not in st.session_state:
        st.session_state.skipped_listings = []

    if analyze_clicked:
        if not url:
            st.error("⚠️ Please enter a BizBuySell URL")
            return

        if "bizbuysell.com" not in url.lower():
            st.error("⚠️ Please enter a valid BizBuySell URL")
            return

        # Progress indicators
        progress_bar = st.progress(0)
        status_text = st.empty()

        try:
            # Check if it's a direct listing URL or a search page
            is_single_listing = is_direct_listing_url(url)

            if is_single_listing:
                # Direct listing URL - analyze just this one business
                status_text.markdown("🔍 **Fetching business listing...**")

                def single_progress(current, total, message):
                    progress = (current / total) * 0.5 if total > 0 else 0
                    progress_bar.progress(progress)
                    status_text.markdown(f"📊 **{message}**")

                business = scrape_single_listing(
                    url,
                    single_progress,
                    scraper_api_key=scraper_api_key
                )

                if not business:
                    st.error("❌ Failed to fetch business listing. The page may be unavailable.")
                    progress_bar.empty()
                    status_text.empty()
                    return

                businesses = [business]
                st.session_state.skipped_listings = []  # No skipped for single listing
            else:
                # Search page - scrape multiple listings
                status_text.markdown("🔍 **Fetching business listings...**")

                def update_progress(current, total, message):
                    # Map scraping progress to 0%-50% of progress bar
                    progress = (current / total) * 0.5 if total > 0 else 0
                    progress_bar.progress(progress)
                    # Show filtering status with icon
                    if "Filtering" in message:
                        status_text.markdown(f"🔍 **{message}**")
                    else:
                        status_text.markdown(f"📊 **{message}**")

                businesses, skipped_listings = scrape_listings(
                    url,
                    limit,
                    update_progress,
                    scraper_api_key=scraper_api_key,
                    filter_excluded=True  # Filter out dental, legal, medical, restaurants
                )

                # Store skipped listings
                st.session_state.skipped_listings = skipped_listings

                if not businesses:
                    st.error("❌ No business listings found. The page structure may have changed or the URL may be invalid.")
                    progress_bar.empty()
                    status_text.empty()
                    return

            # Analysis phase
            status_text.markdown("🧠 **Analyzing opportunities with AI...**")

            def analysis_progress(current, total, message):
                # Map analysis progress to 50%-100% of progress bar
                progress = 0.5 + (current / total) * 0.5 if total > 0 else 0.5
                progress_bar.progress(progress)
                status_text.markdown(f"🧠 **{message}** ({current}/{total})")

            results = analyze_businesses(businesses, progress_callback=analysis_progress)

            progress_bar.progress(1.0)
            if is_single_listing:
                status_text.markdown(f"✅ **Analysis complete!**")
            else:
                status_text.markdown(f"✅ **Analysis complete! Found {len(results)} opportunities.**")

            # Store results
            st.session_state.results = results

        except Exception as e:
            st.error(f"❌ Error: {str(e)}")
            progress_bar.empty()
            status_text.empty()
            return

    # Display results if available
    if st.session_state.results:
        results = st.session_state.results

        # Summary metrics
        st.markdown("---")
        col1, col2, col3, col4 = st.columns(4)

        with col1:
            st.metric("📊 Analyzed", f"{len(results)} businesses")

        with col2:
            avg_score = sum(r.overall_score for r in results) / len(results)
            st.metric("⭐ Avg Score", f"{avg_score:.1f}/100")

        with col3:
            best = results[0]
            st.metric("🏆 Best ROI", f"{best.metrics.roi_percent:.1f}%" if best.metrics.roi_percent else "N/A")

        with col4:
            prices = [r.business.asking_price for r in results if r.business.asking_price]
            if prices:
                avg_price = sum(prices) / len(prices)
                st.metric("💰 Avg Price", format_currency(avg_price))
            else:
                st.metric("💰 Avg Price", "N/A")

        # Render comparison table
        render_comparison_table(results)

        # Render top picks
        render_top_picks(results)

        # Export option
        st.markdown("---")
        st.markdown("### 📥 Export Data")

        col_csv, col_word = st.columns(2)

        # CSV Export (sanitized for compatibility)
        export_data = []
        for r in results:
            export_data.append({
                "Rank": r.rank,
                "Business Name": sanitize_for_csv(r.business.name),
                "URL": r.business.url,
                "Asking Price": r.business.asking_price,
                "Cash Flow": r.business.cash_flow,
                "Gross Revenue": r.business.gross_revenue,
                "ROI %": r.metrics.roi_percent,
                "Payback Years": r.metrics.payback_years,
                "Score": r.overall_score,
                "Strengths": sanitize_for_csv("; ".join(r.swot.strengths)),
                "Weaknesses": sanitize_for_csv("; ".join(r.swot.weaknesses)),
                "Opportunities": sanitize_for_csv("; ".join(r.swot.opportunities)),
                "Threats": sanitize_for_csv("; ".join(r.swot.threats)),
                "Recommendation": sanitize_for_csv(r.recommendation),
            })

        export_df = pd.DataFrame(export_data)
        csv = export_df.to_csv(index=False)

        with col_csv:
            st.download_button(
                label="📥 Download CSV",
                data=csv,
                file_name="business_analysis_report.csv",
                mime="text/csv",
                use_container_width=True
            )

        # Word Export
        with col_word:
            if DOCX_AVAILABLE:
                word_doc = generate_word_report(results)
                if word_doc:
                    st.download_button(
                        label="📄 Download Word Report",
                        data=word_doc,
                        file_name="business_analysis_report.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )
            else:
                st.caption("Word export requires python-docx package")

        # Display skipped listings at the bottom
        if st.session_state.skipped_listings:
            st.markdown("---")
            with st.expander(f"⏭️ Skipped Listings ({len(st.session_state.skipped_listings)} total)", expanded=False):
                st.markdown("These listings were skipped during analysis:")

                # Group by reason
                from collections import defaultdict
                by_reason = defaultdict(list)
                for skip in st.session_state.skipped_listings:
                    by_reason[skip.reason].append(skip)

                for reason, listings in sorted(by_reason.items()):
                    st.markdown(f"**{reason}** ({len(listings)} listings)")
                    for skip in listings:
                        # Truncate name for display but keep full URL as link
                        display_name = skip.name[:60] + "..." if len(skip.name) > 60 else skip.name
                        st.markdown(f"- [{display_name}]({skip.url})")
                    st.markdown("")  # Add spacing between groups


if __name__ == "__main__":
    main()
